# =========================================================
# 🚀 PREMIUM CUSTOMER CHURN DASHBOARD
# =========================================================

import gradio as gr
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import joblib

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.ensemble import GradientBoostingClassifier
from sklearn.metrics import accuracy_score

# =========================================================
# LOAD DATA
# =========================================================

df = pd.read_csv(r"C:\Users\alosh\OneDrive\Desktop\DSDA\Python\files\Python projects\customer_data.csv")

# =========================================================
# DATA PREPROCESSING
# =========================================================

df.columns = df.columns.str.lower()

# Encode categorical columns
label_encoders = {}

for col in ['country', 'gender']:
    le = LabelEncoder()
    df[col] = le.fit_transform(df[col])
    label_encoders[col] = le

# Features and Target
X = df.drop(['customer_id', 'churn'], axis=1)
y = df['churn']

# Train Test Split
X_train, X_test, y_train, y_test = train_test_split(
    X,
    y,
    test_size=0.2,
    random_state=42
)

# Scaling
scaler = StandardScaler()

X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# =========================================================
# MODEL TRAINING
# =========================================================

model = GradientBoostingClassifier()

model.fit(X_train_scaled, y_train)

preds = model.predict(X_test_scaled)

accuracy = accuracy_score(y_test, preds)

# =========================================================
# KPI FUNCTION
# =========================================================

def get_kpis(filtered_df):

    total_customers = len(filtered_df)

    churn_rate = round(
        filtered_df['churn'].mean() * 100,
        2
    )

    avg_balance = round(
        filtered_df['balance'].mean(),
        2
    )

    avg_age = round(
        filtered_df['age'].mean(),
        1
    )

    active_members = filtered_df['active_member'].sum()

    high_risk = len(
        filtered_df[
            (filtered_df['balance'] > 100000) &
            (filtered_df['age'] > 45)
        ]
    )

    html = f"""
    <div style='display:flex; gap:15px; flex-wrap:wrap;'>

        <div style='background:#111827;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>Total Customers</h3>
        <h1>{total_customers}</h1>
        </div>

        <div style='background:#7c3aed;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>Churn Rate</h3>
        <h1>{churn_rate}%</h1>
        </div>

        <div style='background:#059669;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>Avg Balance</h3>
        <h1>${avg_balance}</h1>
        </div>

        <div style='background:#2563eb;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>Average Age</h3>
        <h1>{avg_age}</h1>
        </div>

        <div style='background:#dc2626;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>Active Members</h3>
        <h1>{active_members}</h1>
        </div>

        <div style='background:#ea580c;padding:20px;
        border-radius:15px;width:180px;color:white'>
        <h3>High Risk</h3>
        <h1>{high_risk}</h1>
        </div>

    </div>
    """

    return html

# =========================================================
# CHARTS
# =========================================================

def generate_dashboard(country_filter, gender_filter):

    filtered_df = df.copy()

    if country_filter != "All":
        country_encoded = label_encoders['country'].transform([country_filter])[0]
        filtered_df = filtered_df[
            filtered_df['country'] == country_encoded
        ]

    if gender_filter != "All":
        gender_encoded = label_encoders['gender'].transform([gender_filter])[0]
        filtered_df = filtered_df[
            filtered_df['gender'] == gender_encoded
        ]

    # =====================================================
    # CHART 1 : CHURN DISTRIBUTION
    # =====================================================

    fig1, ax1 = plt.subplots(figsize=(5,4))

    sns.countplot(
        data=filtered_df,
        x='churn',
        palette='viridis',
        ax=ax1
    )

    ax1.set_title("Customer Churn Distribution")

    insight1 = """
    INSIGHT:
    Majority customers are retained but churn count is still significant.

    SUGGESTION:
    Create customer retention campaigns for at-risk customers.
    """

    # =====================================================
    # CHART 2 : BALANCE VS CHURN
    # =====================================================

    fig2, ax2 = plt.subplots(figsize=(5,4))

    sns.boxplot(
        data=filtered_df,
        x='churn',
        y='balance',
        palette='Set2',
        ax=ax2
    )

    ax2.set_title("Balance vs Churn")

    insight2 = """
    INSIGHT:
    High balance customers are more likely to churn.

    SUGGESTION:
    Offer loyalty rewards and premium banking services.
    """

    # =====================================================
    # CHART 3 : AGE VS CHURN
    # =====================================================

    fig3, ax3 = plt.subplots(figsize=(5,4))

    sns.violinplot(
        data=filtered_df,
        x='churn',
        y='age',
        palette='coolwarm',
        ax=ax3
    )

    ax3.set_title("Age Distribution by Churn")

    insight3 = """
    INSIGHT:
    Older customers show higher churn tendency.

    SUGGESTION:
    Improve customer support for older customers.
    """

    # =====================================================
    # CHART 4 : ACTIVE MEMBERS
    # =====================================================

    fig4, ax4 = plt.subplots(figsize=(5,4))

    sns.countplot(
        data=filtered_df,
        x='active_member',
        hue='churn',
        palette='Set1',
        ax=ax4
    )

    ax4.set_title("Active Members vs Churn")

    insight4 = """
    INSIGHT:
    Inactive customers are more likely to churn.

    SUGGESTION:
    Increase engagement campaigns and personalized communication.
    """

    # =====================================================
    # CHART 5 : CREDIT SCORE
    # =====================================================

    fig5, ax5 = plt.subplots(figsize=(5,4))

    sns.histplot(
        data=filtered_df,
        x='credit_score',
        hue='churn',
        kde=True,
        palette='husl',
        ax=ax5
    )

    ax5.set_title("Credit Score Distribution")

    insight5 = """
    INSIGHT:
    Customers with lower credit scores show slightly higher churn.

    SUGGESTION:
    Provide customized financial solutions for low credit score customers.
    """

    return (
        get_kpis(filtered_df),

        fig1, insight1,
        fig2, insight2,
        fig3, insight3,
        fig4, insight4,
        fig5, insight5
    )

# =========================================================
# HIGH RISK CUSTOMERS
# =========================================================

def high_risk_customers():

    risk_df = df[
        (df['balance'] > 100000) &
        (df['age'] > 45)
    ]

    return risk_df.sort_values(
        by='balance',
        ascending=False
    )

# =========================================================
# PREDICTION FUNCTION
# =========================================================

def predict_churn(
    credit_score,
    country,
    gender,
    age,
    tenure,
    balance,
    products_number,
    credit_card,
    active_member,
    estimated_salary
):

    country_encoded = label_encoders['country'].transform([country])[0]

    gender_encoded = label_encoders['gender'].transform([gender])[0]

    input_data = pd.DataFrame([[
        credit_score,
        country_encoded,
        gender_encoded,
        age,
        tenure,
        balance,
        products_number,
        credit_card,
        active_member,
        estimated_salary
    ]], columns=X.columns)

    scaled_data = scaler.transform(input_data)

    prediction = model.predict(scaled_data)[0]

    probability = model.predict_proba(scaled_data)[0][1]

    if prediction == 1:
        result = f"""
        🔴 HIGH CHURN RISK

        Probability: {round(probability*100,2)}%

        Recommendation:
        Provide immediate retention offer and personalized engagement.
        """
    else:
        result = f"""
        🟢 LOW CHURN RISK

        Probability: {round(probability*100,2)}%

        Recommendation:
        Continue maintaining customer satisfaction.
        """

    return result

# =========================================================
# DOWNLOAD REPORT AS DOCX
# =========================================================

from docx import Document
from docx.shared import Inches

def generate_report():

    # Create Word Document
    doc = Document()

    # =====================================================
    # TITLE
    # =====================================================

    doc.add_heading(
        'Customer Churn Analysis Report',
        level=1
    )

    doc.add_paragraph(
        'This report contains customer churn insights, KPIs, '
        'business recommendations, and statistical summary.'
    )

    # =====================================================
    # KPI SECTION
    # =====================================================

    total_customers = len(df)

    churn_rate = round(
        df['churn'].mean() * 100,
        2
    )

    avg_balance = round(
        df['balance'].mean(),
        2
    )

    avg_age = round(
        df['age'].mean(),
        1
    )

    active_members = df['active_member'].sum()

    high_risk = len(
        df[
            (df['balance'] > 100000) &
            (df['age'] > 45)
        ]
    )

    doc.add_heading('Key Performance Indicators', level=2)

    kpi_table = doc.add_table(rows=1, cols=2)

    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    kpis = [
        ('Total Customers', str(total_customers)),
        ('Churn Rate', f'{churn_rate}%'),
        ('Average Balance', f'${avg_balance}'),
        ('Average Age', str(avg_age)),
        ('Active Members', str(active_members)),
        ('High Risk Customers', str(high_risk))
    ]

    for metric, value in kpis:

        row_cells = kpi_table.add_row().cells

        row_cells[0].text = metric
        row_cells[1].text = value

    # =====================================================
    # BUSINESS INSIGHTS
    # =====================================================

    doc.add_heading('Business Insights', level=2)

    insights = [
        (
            'High Balance Customers',
            'Customers with higher balances showed higher churn probability.',
            'Provide loyalty rewards and premium services.'
        ),

        (
            'Older Customers',
            'Older customers demonstrated higher churn tendency.',
            'Improve support and engagement programs.'
        ),

        (
            'Inactive Customers',
            'Inactive members were more likely to churn.',
            'Increase customer engagement campaigns.'
        ),

        (
            'Credit Score Impact',
            'Lower credit score customers showed moderate churn risk.',
            'Provide personalized financial products.'
        )
    ]

    for title, insight, suggestion in insights:

        doc.add_heading(title, level=3)

        doc.add_paragraph(
            f'Insight: {insight}'
        )

        doc.add_paragraph(
            f'Suggestion: {suggestion}'
        )

    # =====================================================
    # DATA SUMMARY
    # =====================================================

    doc.add_heading('Dataset Statistical Summary', level=2)

    summary = df.describe()

    summary_table = doc.add_table(
        rows=1,
        cols=len(summary.columns) + 1
    )

    header_cells = summary_table.rows[0].cells

    header_cells[0].text = 'Statistics'

    for i, col in enumerate(summary.columns):
        header_cells[i + 1].text = col

    for index in summary.index:

        row_cells = summary_table.add_row().cells

        row_cells[0].text = str(index)

        for i, value in enumerate(summary.loc[index]):

            row_cells[i + 1].text = str(round(value, 2))

    # =====================================================
    # FINAL CONCLUSION
    # =====================================================

    doc.add_heading('Final Conclusion', level=2)

    doc.add_paragraph(
        'The machine learning model successfully identified '
        'key churn-driving factors and achieved strong prediction performance. '
        'The organization can use these insights to improve customer retention '
        'and reduce revenue loss.'
    )

    # =====================================================
    # SAVE REPORT
    # =====================================================

    report_path = "Customer_Churn_Report.docx"

    doc.save(report_path)

    return report_path


# =========================================================
# PREMIUM THEME CSS
# =========================================================

css = """
body {
    background-color: #0f172a;
}

.gradio-container {
    background: linear-gradient(to right,#0f172a,#111827);
    color: white;
}

h1,h2,h3 {
    color: white !important;
}

textarea {
    font-size: 15px !important;
}
"""

# =========================================================
# GRADIO DASHBOARD
# =========================================================
with gr.Blocks(css=css, theme=gr.themes.Soft()) as demo:

    gr.Markdown(
        """
        # 🚀 PREMIUM CUSTOMER CHURN DASHBOARD
        
        ### AI-Powered Customer Retention Analytics
        """
    )

    with gr.Tabs():

        # =================================================
        # PAGE 1 : INSIGHTS
        # =================================================

        with gr.Tab("📊 Insights"):

            with gr.Row():

                country_filter = gr.Dropdown(
                    choices=["All"] + list(label_encoders['country'].classes_),
                    value="All",
                    label="Country Filter"
                )

                gender_filter = gr.Dropdown(
                    choices=["All"] + list(label_encoders['gender'].classes_),
                    value="All",
                    label="Gender Filter"
                )

            refresh_btn = gr.Button("Generate Dashboard")

            kpi_output = gr.HTML()

            with gr.Row():
                with gr.Column():
                    chart1 = gr.Plot()
                    insight1 = gr.Textbox(label="Insight & Suggestion")
                with gr.Column():
                    chart2 = gr.Plot()
                    insight2 = gr.Textbox(label="Insight & Suggestion")

            with gr.Row():
                with gr.Column():
                    chart3 = gr.Plot()
                    insight3 = gr.Textbox(label="Insight & Suggestion")
                with gr.Column():
                    chart4 = gr.Plot()
                    insight4 = gr.Textbox(label="Insight & Suggestion")
            
            with gr.Row():
                with gr.Column():
                    chart5 = gr.Plot()
                    insight5 = gr.Textbox(label="Insight & Suggestion")

            refresh_btn.click(
                generate_dashboard,
                inputs=[country_filter, gender_filter],
                outputs=[
                    kpi_output,

                    chart1, insight1,
                    chart2, insight2,
                    chart3, insight3,
                    chart4, insight4,
                    chart5, insight5
                ]
            )

        # =================================================
        # PAGE 2 : HIGH RISK CUSTOMERS
        # =================================================

        with gr.Tab("⚠️ High Risk Customers"):

            gr.Markdown(
                """
                ## High Risk Customers
                
                Customers with:
                - High Balance
                - Higher Age
                - Higher Churn Risk
                """
            )

            risk_table = gr.Dataframe()

            load_btn = gr.Button("Load High Risk Customers")

            load_btn.click(
                high_risk_customers,
                outputs=risk_table
            )

        # =================================================
        # PAGE 3 : PREDICT
        # =================================================

        with gr.Tab("🤖 Predict Churn"):

            gr.Markdown(
                """
                ## Enter Customer Details
                """
            )

            with gr.Row():

                credit_score = gr.Number(label="Credit Score")
                country = gr.Dropdown(
                    choices=list(label_encoders['country'].classes_),
                    label="Country"
                )

                gender = gr.Dropdown(
                    choices=list(label_encoders['gender'].classes_),
                    label="Gender"
                )

            with gr.Row():

                age = gr.Number(label="Age")
                tenure = gr.Number(label="Tenure")
                balance = gr.Number(label="Balance")

            with gr.Row():

                products_number = gr.Number(label="Products Number")
                credit_card = gr.Radio([0,1], label="Has Credit Card")
                active_member = gr.Radio([0,1], label="Active Member")

            estimated_salary = gr.Number(
                label="Estimated Salary"
            )

            predict_btn = gr.Button("Predict Churn")

            prediction_output = gr.Textbox(
                label="Prediction Result",
                lines=6
            )

            predict_btn.click(
                predict_churn,
                inputs=[
                    credit_score,
                    country,
                    gender,
                    age,
                    tenure,
                    balance,
                    products_number,
                    credit_card,
                    active_member,
                    estimated_salary
                ],
                outputs=prediction_output
            )

        # =================================================
        # PAGE 4 : DOWNLOAD REPORT
        # =================================================

        with gr.Tab("📥 Download Report"):

            gr.Markdown(
                """
                ## Download Customer Analytics Report
                """
            )

            download_btn = gr.Button("Generate Report")

            report_file = gr.File()

            download_btn.click(
                generate_report,
                outputs=report_file
            )

# =========================================================
# LAUNCH
# =========================================================

demo.launch(share=True)